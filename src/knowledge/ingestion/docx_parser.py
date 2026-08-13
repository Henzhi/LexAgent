"""文档解析管道 — Word 解析器。

基于 python-docx，提取段落文本。
"""
from __future__ import annotations

import io
from pathlib import Path


class DocxParser:
    """Word 文档解析器

    用法:
        parser = DocxParser()
        text = parser.parse("path/to/document.docx")
    """

    def parse(self, file_path: str) -> str:
        """解析 DOCX 文件"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: uv add python-docx")

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc = Document(str(path))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                # 保留段落间的空行（作为分隔）
                paragraphs.append("")
                continue

            # 如果是标题样式，加标记
            if para.style.name.startswith("Heading"):
                paragraphs.append(f"\n## {text}")
            else:
                paragraphs.append(text)

        return "\n".join(paragraphs)

    def parse_bytes(self, content: bytes) -> str:
        """从字节流解析 DOCX"""
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                paragraphs.append("")
                continue
            if para.style.name.startswith("Heading"):
                paragraphs.append(f"\n## {text}")
            else:
                paragraphs.append(text)
        return "\n".join(paragraphs)
